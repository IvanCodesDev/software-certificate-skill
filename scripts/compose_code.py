#!/usr/bin/env python3
"""Compose provenance-traceable source identification materials."""

from __future__ import annotations

import argparse
import fnmatch
import unicodedata
from pathlib import Path

from common import load_json, now_iso, relative_posix, safe_text, save_json, sha256_file

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt
except ImportError as exc:
    raise SystemExit("python-docx is required; use the bundled document runtime.") from exc


# A4 filing geometry shared by line wrapping and DOCX writing.
PAGE_WIDTH_MM = 210.0
PAGE_HEIGHT_MM = 297.0
MARGIN_TOP_MM = 22.0
MARGIN_BOTTOM_MM = 25.4
MARGIN_LEFT_MM = 20.0
MARGIN_RIGHT_MM = 20.0
MM_TO_PT = 72.0 / 25.4
CONTENT_WIDTH_PT = (PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM) * MM_TO_PT
CONTENT_HEIGHT_PT = (PAGE_HEIGHT_MM - MARGIN_TOP_MM - MARGIN_BOTTOM_MM) * MM_TO_PT
ASCII_FONT = "Consolas"
EAST_ASIA_FONT = "SimSun"
# Widest plausible monospace advance (em) across Consolas and its substitutes
# (Liberation Mono / DejaVu Sans Mono / Courier family on LibreOffice), so a
# pre-wrapped line never re-wraps in any rendering engine.
ASCII_COLUMN_EM = 0.61


def exact_line_spacing_pt(lines_per_page: int) -> float:
    """Exact leading that lets lines_per_page lines fill the page body.

    Computed in whole twips with ~2pt slack per page so Word and LibreOffice
    rounding can never push the last line onto the next page.
    """
    content_twips = int(CONTENT_HEIGHT_PT * 20)
    return ((content_twips - 40) // lines_per_page) / 20.0


def max_columns_for(font_size: float, requested: int) -> int:
    fit = int(CONTENT_WIDTH_PT / (ASCII_COLUMN_EM * font_size))
    return max(20, min(requested, fit))


DEFAULT_EXTENSIONS = {
    ".py", ".cs", ".java", ".kt", ".js", ".jsx", ".ts", ".tsx", ".vue",
    ".go", ".rs", ".c", ".h", ".cpp", ".cc", ".php", ".rb", ".swift", ".sql",
    ".cshtml", ".razor", ".html"
}
SENSITIVE_TOKENS = ("privatekey", "private_key", "id_rsa", "id_dsa", "secret.key", ".env")


def matches_any(relative: str, patterns: list[str]) -> bool:
    normalized = relative.replace("\\", "/")
    return any(fnmatch.fnmatch(normalized, pattern) or Path(normalized).match(pattern) for pattern in patterns)


def select_files(root: Path, manifest: dict) -> list[Path]:
    ordered = manifest.get("ordered_files", [])
    exclude = list(manifest.get("exclude", []))
    if ordered:
        candidates = [(root / item).resolve() for item in ordered]
    else:
        candidates_set: set[Path] = set()
        for pattern in manifest.get("include", ["**/*"]):
            candidates_set.update(path.resolve() for path in root.glob(pattern) if path.is_file())
        candidates = sorted(candidates_set, key=lambda p: relative_posix(p, root).lower())
    result: list[Path] = []
    for path in candidates:
        try:
            rel = relative_posix(path, root)
        except ValueError as exc:
            raise ValueError(f"Manifest path escapes project root: {path}") from exc
        if not path.is_file() or path.suffix.lower() not in DEFAULT_EXTENSIONS:
            continue
        if matches_any(rel, exclude):
            continue
        if any(token in rel.lower() for token in SENSITIVE_TOKENS):
            continue
        result.append(path)
    return result


def char_columns(character: str) -> int:
    """Display columns of one character in a CJK-capable monospace grid."""
    return 2 if unicodedata.east_asian_width(character) in ("W", "F", "A") else 1


def wrap_source_line(line: str, max_columns: int) -> list[str]:
    expanded = line.expandtabs(4).rstrip()
    if not expanded:
        return [""]
    segments: list[str] = []
    current: list[str] = []
    width = 0
    for character in expanded:
        columns = char_columns(character)
        if current and width + columns > max_columns:
            segments.append("".join(current))
            current, width = [], 0
        current.append(character)
        width += columns
    segments.append("".join(current))
    return segments


def collect_lines(root: Path, files: list[Path], max_chars: int) -> tuple[list[str], list[dict], list[dict], int]:
    """Build the filing corpus: every material line is a non-empty code line.

    Reviewers read "每页不少于 50 行" as 50 lines of effective code, so blank
    lines never enter the corpus — a page padded with blanks is the most
    common reason code materials get returned for correction.
    """
    output: list[str] = []
    mapping: list[dict] = []
    file_records: list[dict] = []
    blank_lines = 0
    for path in files:
        text = safe_text(path, max_bytes=20_000_000)
        if text is None:
            continue
        rel = relative_posix(path, root)
        digest = sha256_file(path)
        start = len(output) + 1
        original_lines = text.splitlines()
        for original_no, original in enumerate(original_lines, 1):
            if not original.strip():
                blank_lines += 1
                continue
            wrapped = wrap_source_line(original, max_chars)
            for segment_no, segment in enumerate(wrapped, 1):
                output.append(segment)
                mapping.append({
                    "output_line": len(output),
                    "file": rel,
                    "file_sha256": digest,
                    "source_line": original_no,
                    "segment": segment_no
                })
        file_records.append({
            "path": rel,
            "sha256": digest,
            "original_lines": len(original_lines),
            "output_start_line": start,
            "output_end_line": len(output)
        })
    return output, mapping, file_records, blank_lines


def group_logical_lines(lines: list[str], mapping: list[dict]) -> list[list[str]]:
    """Group rendered rows into logical source lines (a wrap stays together)."""
    groups: list[list[str]] = []
    for text, meta in zip(lines, mapping):
        if meta["segment"] == 1:
            groups.append([])
        groups[-1].append(text)
    return groups


def paginate(lines: list[str], mapping: list[dict], per_page: int) -> tuple[list[list[str]], int]:
    """Fill each page to a shared row grid, never below `per_page` code lines.

    Reviewers count effective source lines, so a long line wrapped over two
    rows is still one line: paginating on rendered rows leaves most pages
    short of the quota. The grid is sized so the page needing the most rows
    still fits, then every page is packed full — pages without wraps simply
    carry a few extra code lines instead of trailing blank space.
    """
    groups = group_logical_lines(lines, mapping)
    if not groups:
        return [], per_page
    # Widest row count any `per_page` run of code lines can need.
    grid = max(
        sum(len(group) for group in groups[start:start + per_page])
        for start in range(0, max(1, len(groups) - per_page + 1))
    )
    for _ in range(64):
        pages: list[list[str]] = []
        counts: list[int] = []
        current: list[str] = []
        logical = 0
        for group in groups:
            if current and len(current) + len(group) > grid:
                pages.append(current)
                counts.append(logical)
                current, logical = [], 0
            current.extend(group)
            logical += 1
        if current:
            pages.append(current)
            counts.append(logical)
        if all(count >= per_page for count in counts[:-1]):
            return pages, grid
        grid += 1
    return pages, grid


def logical_line_count(page: list[str], mapping: list[dict], start_index: int) -> int:
    return sum(1 for offset in range(len(page)) if mapping[start_index + offset]["segment"] == 1)


def write_txt(path: Path, pages: list[list[str]]) -> None:
    path.write_text("\n\f\n".join("\n".join(page) for page in pages) + "\n", encoding="utf-8")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None:
            cols.addprevious(pg_num_type)
        else:
            sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def write_docx(path: Path, pages: list[list[str]], facts: dict, rows_per_page: int,
               font_size: float = 9.0, line_spacing: float | None = None,
               page_number_start: int = 1, filing_total: int | None = None) -> None:
    grid = exact_line_spacing_pt(rows_per_page)
    spacing = min(line_spacing, grid) if line_spacing else grid
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(17.5)
    set_page_number_start(section, page_number_start)
    normal = doc.styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(font_size)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(spacing)
    normal.paragraph_format.widow_control = False

    # Unified filing header: full software name + version on the left, the
    # page number on the right, nothing in the footer. The template's Header
    # style ships a centre tab stop that OOXML would merge with ours and
    # capture the tab character, so it must be stripped first.
    header_style_ppr = doc.styles["Header"].element.pPr
    if header_style_ppr is not None:
        style_tabs = header_style_ppr.find(qn("w:tabs"))
        if style_tabs is not None:
            header_style_ppr.remove(style_tabs)
    header = section.header.paragraphs[0]
    header.paragraph_format.tab_stops.add_tab_stop(
        Mm(PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM), WD_TAB_ALIGNMENT.RIGHT)
    header.add_run(f"{facts.get('software_full_name', '软件源程序')}{facts.get('version', '')}")
    header.add_run("\t第 ")
    add_page_field(header)
    header.add_run(" 页")
    for run in header.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        run.font.size = Pt(9)

    # pageBreakBefore keeps every physical page at exactly lines_per_page
    # paragraphs; an explicit break paragraph would consume one line slot.
    for page_index, lines in enumerate(pages):
        for line_index, line in enumerate(lines):
            paragraph = doc.add_paragraph()
            if page_index and not line_index:
                paragraph.paragraph_format.page_break_before = True
            paragraph.add_run(line if line else " ")
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project", required=True, type=Path)
    parser.add_argument("--manifest", required=True, type=Path)
    parser.add_argument("--facts", type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--lines-per-page", type=int, default=50,
                        help="Effective (logical) source lines per page")
    parser.add_argument("--max-chars", type=int, default=87)
    parser.add_argument("--font-size", type=float, default=9.0)
    parser.add_argument("--line-spacing", type=float, default=None,
                        help="Exact leading in pt; defaults to page body height / rendered rows")
    args = parser.parse_args()
    if args.lines_per_page < 50:
        parser.error("--lines-per-page must be at least 50")
    root = args.project.resolve()
    manifest = load_json(args.manifest)
    facts = load_json(args.facts) if args.facts else {}
    files = select_files(root, manifest)
    if not files:
        parser.error("No eligible source files selected by the manifest")
    max_columns = max_columns_for(args.font_size, args.max_chars)
    lines, mapping, file_records, blank_lines = collect_lines(root, files, max_columns)
    if not lines:
        parser.error("Selected files contain no readable source lines")
    full_pages, rows_per_page = paginate(lines, mapping, args.lines_per_page)
    total_pages = len(full_pages)
    grid_spacing = exact_line_spacing_pt(rows_per_page)
    line_spacing = min(args.line_spacing, grid_spacing) if args.line_spacing else grid_spacing
    if total_pages < 60:
        groups = {"all": list(range(total_pages))}
        selection = "all_under_60_pages"
    else:
        groups = {"front_30": list(range(30)), "back_30": list(range(total_pages - 30, total_pages))}
        selection = "first_30_and_last_30_separate_volumes"
    # Pages hold a variable number of rendered rows, so line offsets must be
    # accumulated rather than derived from a fixed page size.
    page_offsets: list[int] = []
    running = 0
    for page in full_pages:
        page_offsets.append(running)
        running += len(page)
    output = args.output_dir.resolve()
    output.mkdir(parents=True, exist_ok=True)
    full_txt = output / "source-full.txt"
    full_docx = output / "source-full.docx"
    write_txt(full_txt, full_pages)
    write_docx(full_docx, full_pages, facts, rows_per_page, args.font_size, line_spacing)
    artifacts = {
        "full_txt": {"path": str(full_txt), "sha256": sha256_file(full_txt)},
        "full_docx": {"path": str(full_docx), "sha256": sha256_file(full_docx)},
    }
    filing_total = sum(len(indices) for indices in groups.values())
    filing_groups = {}
    for key, indices in groups.items():
        pages = [full_pages[index] for index in indices]
        source_numbers = [index + 1 for index in indices]
        txt_path = output / f"source-{key.replace('_', '-')}.txt"
        docx_path = output / f"source-{key.replace('_', '-')}.docx"
        write_txt(txt_path, pages)
        # Filing volumes carry one continuous page sequence: front 1-30 and
        # back 31-60 when split, otherwise 1-N for the single volume.
        write_docx(docx_path, pages, facts, rows_per_page, args.font_size, line_spacing,
                   page_number_start=31 if key == "back_30" else 1, filing_total=filing_total)
        artifacts[f"{key}_txt"] = {"path": str(txt_path), "sha256": sha256_file(txt_path)}
        artifacts[f"{key}_docx"] = {"path": str(docx_path), "sha256": sha256_file(docx_path)}
        filing_groups[key] = {
            "logical_source_pages": source_numbers,
            "page_count": len(pages),
            "first_output_line": page_offsets[indices[0]] + 1 if indices else None,
            "last_output_line": page_offsets[indices[-1]] + len(pages[-1]) if indices else None,
        }

    line_pages = []
    for page_no, page_lines in enumerate(full_pages, 1):
        offset = page_offsets[page_no - 1]
        line_pages.append({
            "page": page_no,
            "start_output_line": offset + 1,
            "end_output_line": offset + len(page_lines),
            "line_count": len(page_lines),
            "effective_lines": logical_line_count(page_lines, mapping, offset),
        })
    provenance = {
        "schema_version": "1.0",
        "generated_at": now_iso(),
        "project_root": str(root),
        "manifest": str(args.manifest.resolve()),
        "manifest_review": manifest.get("review", {}),
        "ordered_files_confirmed": bool(manifest.get("ordered_files")),
        "selection_policy": manifest.get("selection_policy", {}),
        "file_decisions": manifest.get("file_decisions", []),
        "lines_per_page": args.lines_per_page,
        "rows_per_page": rows_per_page,
        "max_chars": max_columns,
        "requested_max_chars": args.max_chars,
        "original_line_count": sum(item.get("original_lines", 0) for item in file_records),
        "blank_lines_excluded": blank_lines,
        "full_line_count": len(lines),
        "full_page_count": total_pages,
        "filing_page_count": sum(item["page_count"] for item in filing_groups.values()),
        "selection": selection,
        "filing_groups": filing_groups,
        "files": file_records,
        "pages": line_pages,
        "line_mapping": mapping,
        "layout": {"font_ascii": ASCII_FONT, "font_east_asia": EAST_ASIA_FONT,
                   "font_size_pt": args.font_size, "line_spacing_pt": line_spacing,
                   "max_display_columns": max_columns, "cjk_columns": 2,
                   "pagination": "page_break_before",
                   "line_policy": "blank lines excluded; pages cut on effective source lines",
                   "rows_per_page": rows_per_page,
                   "header": "software name + version left, page number right, empty footer",
                   "page_numbering": "front 1-30 and back 31-60 when split, otherwise 1-N"},
        "artifacts": artifacts,
    }
    provenance_path = output / "source-provenance.json"
    save_json(provenance_path, provenance)
    print(f"SOURCE_OUTPUT={output}")
    print(f"FILES={len(file_records)} LINES={len(lines)} FULL_PAGES={total_pages} FILING_PAGES={sum(item['page_count'] for item in filing_groups.values())} SELECTION={selection}")
    print(f"PROVENANCE={provenance_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
