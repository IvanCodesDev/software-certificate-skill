#!/usr/bin/env python3
"""Build a polished, evidence-led A4 DOCX with a real Word TOC field."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from common import load_json, sha256_file

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING,
                                WD_TAB_ALIGNMENT, WD_TAB_LEADER)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("python-docx is required; use the bundled document runtime.") from exc


def color(value: str) -> RGBColor:
    return RGBColor.from_string(value.lstrip("#"))


def set_east_asia(run, font_name: str) -> None:
    """Set all OOXML font slots so numbers and Chinese never split families."""
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{key}"), None)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), font_name)


def set_style_fonts(style, chinese: str, latin: str | None = None) -> None:
    latin = latin or chinese
    style.font.name = latin
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{key}"), None)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), chinese)
    fonts.set(qn("w:cs"), latin)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color_hex: str, size: int = 5) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color_hex)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_paragraph_bottom_border(paragraph, color_hex: str, size: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex)


def keep_table_row(row, allow_split: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if not allow_split and tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_outline_level(style, level: int) -> None:
    p_pr = style.element.get_or_add_pPr()
    node = p_pr.find(qn("w:outlineLvl"))
    if node is None:
        node = OxmlElement("w:outlineLvl")
        p_pr.append(node)
    node.set(qn("w:val"), str(level))


def ensure_custom_style(doc: Document, name: str, base: str, size: float, font: str,
                        color_hex: str, bold: bool = False):
    styles = doc.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base]
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color(color_hex)
    set_style_fonts(style, font)
    return style


def configure_styles(doc: Document, theme: dict) -> None:
    fonts, sizes, colors = theme["fonts"], theme["sizes"], theme["colors"]
    normal = doc.styles["Normal"]
    set_style_fonts(normal, fonts["body_cn"], fonts["latin"])
    normal.font.size = Pt(sizes["body"])
    normal.font.color.rgb = color(colors["text"])
    pf = normal.paragraph_format
    pf.line_spacing = 1.45
    pf.space_after = Pt(4)
    pf.first_line_indent = Mm(7.4)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True

    heading_specs = [("Heading 1", sizes["h1"], 0), ("Heading 2", sizes["h2"], 1), ("Heading 3", sizes["h3"], 2)]
    for name, size, level in heading_specs:
        style = doc.styles[name]
        set_style_fonts(style, fonts["heading_cn"], fonts.get("heading_latin", fonts["heading_cn"]))
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color(colors["text"])
        style.paragraph_format.space_before = Pt(10 if level == 0 else 7)
        style.paragraph_format.space_after = Pt(8 if level == 0 else 5)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True
        set_outline_level(style, level)

    ensure_custom_style(doc, "Manual Eyebrow", "Normal", sizes["small"], fonts["latin"], colors["secondary"], True)
    lead = ensure_custom_style(doc, "Manual Lead", "Normal", sizes["body"], fonts["body_cn"], colors["text"])
    set_style_fonts(lead, fonts["body_cn"], fonts["latin"])
    lead.paragraph_format.first_line_indent = Pt(0)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.4
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption = ensure_custom_style(doc, "Manual Caption", "Normal", sizes["caption"], fonts["body_cn"], colors["text"])
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    evidence = ensure_custom_style(doc, "Evidence Line", "Normal", sizes["small"], fonts["mono"], colors["secondary"])
    evidence.paragraph_format.space_before = Pt(4)
    toc_title = ensure_custom_style(doc, "Manual Toc Title", "Normal", sizes["h1"], fonts["heading_cn"], colors["primary"], True)
    toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.first_line_indent = Pt(0)
    toc_title.paragraph_format.space_before = Pt(8)
    toc_title.paragraph_format.space_after = Pt(8)


def configure_page(doc: Document, theme: dict) -> None:
    page = theme["page"]
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(page["top_mm"])
        section.bottom_margin = Mm(page["bottom_mm"])
        section.left_margin = Mm(page["inner_mm"])
        section.right_margin = Mm(page["outer_mm"])
        section.gutter = Mm(page.get("gutter_mm", 0))
        section.header_distance = Mm(page.get("header_mm", 10))
        section.footer_distance = Mm(page.get("footer_mm", 10))
        section.different_first_page_header_footer = True


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def configure_header_footer(doc: Document, facts: dict, theme: dict) -> None:
    colors, fonts, sizes = theme["colors"], theme["fonts"], theme["sizes"]
    label = facts.get("software_short_name") or facts.get("software_full_name", "软件操作手册")
    version = facts.get("version", "")
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(label)
        set_east_asia(r, fonts["body_cn"])
        r.font.size = Pt(sizes["small"])
        r.font.color.rgb = color(colors["muted"])
        r = p.add_run("\t" + version)
        set_east_asia(r, fonts["latin"])
        r.font.size = Pt(sizes["small"])
        r.font.color.rgb = color(colors["muted"])
        set_paragraph_bottom_border(p, colors["rule"], 4)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("第 ")
        set_east_asia(r, fonts["body_cn"])
        r.font.size = Pt(sizes["small"])
        r.font.color.rgb = color(colors["muted"])
        add_field(p, " PAGE ", "1")
        r = p.add_run(" 页")
        set_east_asia(r, fonts["body_cn"])
        r.font.size = Pt(sizes["small"])
        r.font.color.rgb = color(colors["muted"])


def add_cover(doc: Document, page: dict, facts: dict, document: dict, theme: dict) -> None:
    colors, fonts, sizes = theme["colors"], theme["fonts"], theme["sizes"]
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(page.get("title") or facts.get("software_full_name", "软件操作手册"))
    set_east_asia(run, fonts["title_cn"])
    run.font.size = Pt(sizes["cover_title"])
    run.font.bold = True
    run.font.color.rgb = color(colors["text"])

    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p.paragraph_format.space_after = Pt(30)
    run = version_p.add_run(f"{facts.get('version', '')} 版")
    set_east_asia(run, fonts["body_cn"])
    run.font.size = Pt(sizes["cover_subtitle"])
    run.font.color.rgb = color(colors["text"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(document.get("title", "操作手册"))
    set_east_asia(run, fonts["heading_cn"])
    run.font.size = Pt(sizes.get("cover_document", 18))
    run.font.bold = True
    run.font.color.rgb = color(colors["text"])

    for _ in range(8):
        doc.add_paragraph()
    holder = doc.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = holder.add_run(str(facts.get("rightsholder", "")))
    set_east_asia(run, fonts["body_cn"])
    run.font.size = Pt(sizes["body"])
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run(str(document.get("date", "")))
    set_east_asia(run, fonts["body_cn"])
    run.font.size = Pt(sizes["body"])


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p_pr = paragraph._p.find(qn("w:pPr"))
    insert_at = 1 if p_pr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text: str, anchor: str, theme: dict, bold: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), theme["fonts"]["latin"])
    fonts.set(qn("w:hAnsi"), theme["fonts"]["latin"])
    fonts.set(qn("w:eastAsia"), theme["fonts"]["body_cn"])
    r_pr.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), "000000")
    r_pr.append(color_node)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def append_field_start(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    for node in (begin, instr, separate):
        run._r.append(node)


def append_field_end(paragraph) -> None:
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_toc(doc: Document, page: dict, theme: dict, entries: list[dict[str, Any]]) -> None:
    doc.add_paragraph(page.get("title", "目录"), style="Manual Toc Title")
    if not entries:
        toc = doc.add_paragraph()
        add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ', "更新目录")
        return
    for index, entry in enumerate(entries):
        toc = doc.add_paragraph()
        toc.paragraph_format.space_before = Pt(2 if index else 6)
        toc.paragraph_format.space_after = Pt(2)
        if int(entry["level"]) > 1:
            toc.paragraph_format.left_indent = Mm(6 * (int(entry["level"]) - 1))
        toc.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        if index == 0:
            append_field_start(toc, ' TOC \\o "1-3" \\h \\z \\u ')
        add_hyperlink(toc, str(entry["title"]), str(entry["bookmark"]), theme,
                      bold=int(entry["level"]) == 1)
        toc.add_run("\t")
        add_field(toc, f" PAGEREF {entry['bookmark']} \\h ", str(index + 1))
        if index == len(entries) - 1:
            append_field_end(toc)


def add_evidence_line(doc: Document, evidence_ids: list[str], theme: dict) -> None:
    if not evidence_ids or not theme.get("features", {}).get("show_evidence_ids", False):
        return
    p = doc.add_paragraph(style="Evidence Line")
    p.add_run("EVIDENCE  ·  " + "  /  ".join(evidence_ids))


def format_cell(cell, theme: dict, *, bold: bool = False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT, size: float | None = None) -> None:
    fonts, sizes = theme["fonts"], theme["sizes"]
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.25
        for run in paragraph.runs:
            set_east_asia(run, fonts["heading_cn"] if bold else fonts["body_cn"])
            run.font.size = Pt(size or sizes.get("table", 10.5))
            run.font.bold = bold
            run.font.color.rgb = color(theme["colors"]["text"])


DXA_PER_MM = 56.6929


def set_table_grid(table, column_mm: list[float]) -> None:
    """Write explicit tblW and tblGrid widths.

    python-docx only writes per-cell tcW; with the table-level width left at
    "auto", WPS stretches fixed-layout tables to the full text column while
    Word/LibreOffice honour the 160mm grid, so every engine must be pinned to
    the same explicit numbers.
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        anchor = tbl_pr.find(qn("w:jc"))
        if anchor is None:
            anchor = tbl_pr.find(qn("w:tblLayout"))
        if anchor is not None:
            anchor.addprevious(tbl_w)
        else:
            tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(round(sum(column_mm) * DXA_PER_MM))))
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for column, mm in zip(grid.findall(qn("w:gridCol")), column_mm):
            column.set(qn("w:w"), str(int(round(mm * DXA_PER_MM))))


def add_table(doc: Document, block: dict, theme: dict) -> None:
    colors, fonts = theme["colors"], theme["fonts"]
    headers = [str(v) for v in block.get("headers", [])]
    rows = block.get("rows", [])
    if not headers:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = block.get("widths_mm")
    if not widths:
        widths = [38, 122] if len(headers) == 2 else [160 / len(headers)] * len(headers)
    set_table_grid(table, [float(widths[idx] if idx < len(widths) else widths[-1])
                           for idx in range(len(headers))])
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], colors["panel"])
        set_cell_border(header_cells[idx], colors["rule"], 4)
        set_cell_margins(header_cells[idx])
        header_cells[idx].width = Mm(float(widths[idx] if idx < len(widths) else widths[-1]))
        format_cell(header_cells[idx], theme, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    keep_table_row(table.rows[0])
    # Serial-number style columns read better centered than left-aligned.
    centered_columns = {
        idx for idx, header in enumerate(headers)
        if str(header).strip() in {"序号", "编号", "页码", "图号"}
        or (rows and all(re.fullmatch(r"\d{1,4}", str(row[idx]).strip())
                         for row in rows if idx < len(row)))
    }
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            cell.text = str(values[idx]) if idx < len(values) else ""
            set_cell_border(cell, colors["rule"], 4)
            set_cell_margins(cell)
            cell.width = Mm(float(widths[idx] if idx < len(widths) else widths[-1]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            format_cell(cell, theme, alignment=WD_ALIGN_PARAGRAPH.CENTER if idx in centered_columns
                        else WD_ALIGN_PARAGRAPH.LEFT)
        keep_table_row(table.rows[-1])
    # Do not append a spacer paragraph after a table.  When the table exactly
    # fills a page, Word/LibreOffice can push that otherwise-empty paragraph to
    # a new page and create a header/footer-only trailing sheet.


def add_note(doc: Document, block: dict, theme: dict) -> None:
    """Render callouts as a self-contained shaded paragraph.

    A single-cell table box fuses visually with an adjacent data table in
    Word/WPS (two sibling tables render without a gap) and their widths are
    resolved by different layout rules, which reads as a misaligned box. A
    bordered paragraph always stands on its own line with its own spacing,
    and its indents align it with the 160mm table grid.
    """
    colors, fonts = theme["colors"], theme["fonts"]
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    # Word/WPS draw the callout edge at (indent - border space), measured on
    # real renders. indent 5.32mm - space 8pt(2.82mm) puts the box flush with
    # the 160mm table grid while keeping a 2.8mm text inset on each side.
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "8")
        edge.set(qn("w:color"), colors["rule"])
        borders.append(edge)
    p_pr.append(borders)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), colors["soft"])
    p_pr.append(shading)
    fmt = p.paragraph_format
    fmt.left_indent = Mm(5.32)
    fmt.right_indent = Mm(5.32)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.3
    if block.get("title"):
        r = p.add_run(str(block["title"]) + "  ")
        set_east_asia(r, fonts["heading_cn"])
        r.font.bold = True
        r.font.color.rgb = color(colors["secondary"])
    r = p.add_run(str(block.get("text", "")))
    set_east_asia(r, fonts["body_cn"])


def add_facts(doc: Document, block: dict, theme: dict) -> None:
    items = [item for item in block.get("items", []) if item.get("label") or item.get("value")]
    if not items:
        return
    colors = theme["colors"]
    columns = min(3, len(items))
    for start in range(0, len(items), columns):
        group = items[start:start + columns]
        table = doc.add_table(rows=2, cols=len(group))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        width = 160 / len(group)
        set_table_grid(table, [width] * len(group))
        for index, item in enumerate(group):
            label_cell, value_cell = table.cell(0, index), table.cell(1, index)
            label_cell.text = str(item.get("label", ""))
            value_cell.text = str(item.get("value", ""))
            for cell in (label_cell, value_cell):
                cell.width = Mm(width)
                set_cell_border(cell, colors["rule"], 4)
                set_cell_margins(cell, 100, 140, 100, 140)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(label_cell, colors["soft"])
            format_cell(label_cell, theme, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
            format_cell(value_cell, theme, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
        # Avoid empty inter-table paragraphs.  They are semantically inert and
        # may spill onto a separate page during LibreOffice pagination.


def add_subheading(doc: Document, block: dict, theme: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    keep_with_next(p)
    run = p.add_run(str(block.get("text", "")))
    set_east_asia(run, theme["fonts"]["heading_cn"])
    run.font.size = Pt(theme["sizes"].get("body", 12))
    run.font.bold = True


def add_placeholder_image(doc: Document, block: dict, theme: dict) -> None:
    colors, fonts = theme["colors"], theme["fonts"]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_grid(table, [160])
    cell = table.cell(0, 0)
    cell.width = Mm(160)
    set_cell_shading(cell, colors["soft"])
    set_cell_border(cell, colors["rule"], 6)
    table.rows[0].height = Mm(72)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("真实界面截图区域\n\n保持原始宽高比")
    set_east_asia(r, fonts["heading_cn"])
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = color(colors["muted"])
    if block.get("caption"):
        doc.add_paragraph(str(block["caption"]), style="Manual Caption")


def add_image(doc: Document, block: dict, base: Path) -> None:
    image = (base / block["path"]).resolve() if not Path(block["path"]).is_absolute() else Path(block["path"])
    if not image.is_file():
        raise FileNotFoundError(f"Image not found: {image}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Mm(float(block.get("width_mm", 160))))
    if block.get("caption"):
        doc.add_paragraph(str(block["caption"]), style="Manual Caption")


def add_blocks(doc: Document, blocks: list[dict], theme: dict, base: Path) -> None:
    for block in blocks:
        kind = block.get("type", "paragraph")
        if kind == "paragraph":
            doc.add_paragraph(str(block.get("text", "")))
        elif kind == "lead":
            doc.add_paragraph(str(block.get("text", "")), style="Manual Lead")
        elif kind == "steps":
            # Literal numbering is deliberate: Word's built-in List Number style
            # continues across unrelated feature sections and produces misleading
            # 5/9/13... starts.  A hanging-indent paragraph is deterministic in
            # Word, WPS and LibreOffice and restarts every steps block at 1.
            for index, item in enumerate(block.get("items", []), 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Mm(8)
                p.paragraph_format.first_line_indent = Mm(-5.5)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.35
                number = p.add_run(f"{index}. ")
                set_east_asia(number, theme["fonts"]["heading_cn"])
                number.bold = True
                body = p.add_run(str(item))
                set_east_asia(body, theme["fonts"]["body_cn"])
        elif kind == "bullets":
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Mm(8)
                p.paragraph_format.first_line_indent = Mm(-3.5)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(str(item))
                set_east_asia(run, theme["fonts"]["body_cn"])
        elif kind == "subheading":
            add_subheading(doc, block, theme)
        elif kind == "facts":
            add_facts(doc, block, theme)
        elif kind == "note":
            add_note(doc, block, theme)
        elif kind == "table":
            add_table(doc, block, theme)
        elif kind == "image":
            add_image(doc, block, base)
        elif kind == "placeholder_image":
            add_placeholder_image(doc, block, theme)
        elif kind == "code":
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            code_table = doc.add_table(rows=1, cols=1)
            code_table.autofit = False
            set_table_grid(code_table, [160])
            set_cell = code_table.cell(0, 0)
            set_cell.width = Mm(160)
            set_cell_shading(set_cell, theme["colors"]["soft"])
            set_cell_border(set_cell, theme["colors"]["rule"], 4)
            set_cell.text = str(block.get("text", ""))
            for run in set_cell.paragraphs[0].runs:
                set_east_asia(run, theme["fonts"]["mono"])
                run.font.size = Pt(8.5)
        else:
            raise ValueError(f"Unsupported block type: {kind}")


def add_content_page(doc: Document, page: dict, theme: dict, base: Path,
                     bookmark: str | None = None, bookmark_id: int = 0,
                     page_break_before: bool = False) -> None:
    if page.get("eyebrow") and theme.get("features", {}).get("show_eyebrow", False):
        doc.add_paragraph(str(page["eyebrow"]), style="Manual Eyebrow")
    level = int(page.get("level", 1 if page.get("kind") == "section" else 2))
    heading = doc.add_heading(str(page.get("title", "未命名页面")), level=level)
    # Use a heading property instead of a standalone page-break paragraph.
    # When a preceding table or image already fills the page, a standalone
    # break can spill onto its own page and leave only the header/footer.
    heading.paragraph_format.page_break_before = page_break_before
    for run in heading.runs:
        set_east_asia(run, theme["fonts"]["heading_cn"])
    if bookmark:
        add_bookmark(heading, bookmark, bookmark_id)
    keep_with_next(heading)
    if page.get("lead"):
        doc.add_paragraph(str(page["lead"]), style="Manual Lead")
    add_evidence_line(doc, [str(v) for v in page.get("evidence_ids", [])], theme)
    add_blocks(doc, page.get("blocks", []), theme, base)


def build(input_path: Path, theme_path: Path, output_path: Path) -> None:
    payload = load_json(input_path)
    theme = load_json(theme_path)
    facts = payload.get("facts", {})
    document = payload.get("document", {})
    pages = payload.get("pages", [])
    if not pages:
        raise ValueError("manual input contains no pages")
    doc = Document()
    configure_page(doc, theme)
    configure_styles(doc, theme)
    configure_header_footer(doc, facts, theme)
    set_update_fields(doc)
    doc.core_properties.title = f"{facts.get('software_full_name', '')}{document.get('title', '操作手册')}"
    doc.core_properties.subject = document.get("type", "软件著作权鉴别材料")
    doc.core_properties.author = facts.get("rightsholder", "")
    doc.core_properties.comments = f"theme={theme.get('id', '')}; edition={document.get('edition', '')}"

    toc_entries: list[dict[str, Any]] = []
    bookmark_by_index: dict[int, tuple[str, int]] = {}
    bookmark_id = 1
    for index, page in enumerate(pages):
        if page.get("kind") in {"cover", "toc"}:
            continue
        bookmark = f"_SoftCertToc{bookmark_id:04d}"
        level = int(page.get("level", 1 if page.get("kind") == "section" else 2))
        toc_entries.append({"title": page.get("title", "未命名页面"), "level": level, "bookmark": bookmark})
        bookmark_by_index[index] = (bookmark, bookmark_id)
        bookmark_id += 1

    for index, page in enumerate(pages):
        kind = page.get("kind")
        if kind == "cover":
            if index:
                doc.add_page_break()
            add_cover(doc, page, facts, document, theme)
        elif kind == "toc":
            if index:
                doc.add_page_break()
            add_toc(doc, page, theme, toc_entries)
        else:
            bookmark, bookmark_id = bookmark_by_index.get(index, (None, 0))
            previous_kind = pages[index - 1].get("kind") if index else None
            # A hard break is useful after the TOC and before an intentionally
            # separated detail page.  Ordinary chapters flow naturally: this
            # avoids LibreOffice producing a header/footer-only page when the
            # preceding table or screenshot already ended at a page boundary.
            force_break = previous_kind == "toc" or page.get("subpage_kind") == "detail"
            add_content_page(doc, page, theme, input_path.parent, bookmark, bookmark_id,
                             page_break_before=force_break)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"DOCX={output_path.resolve()}")
    print(f"INPUT_PAGES={len(pages)} THEME={theme.get('id')} SHA256={sha256_file(output_path)}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--theme", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build(args.input.resolve(), args.theme.resolve(), args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
