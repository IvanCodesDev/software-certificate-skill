from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from build_manual import add_blocks
from common import sha256_file
from convert_document import ConversionFailure, run_isolated
from generate_application_form import compress
from generate_manual_content import capability_blocks, content_quality, effective_prerequisites
from product_verify import screenshot_release_check


class QualityGateTests(unittest.TestCase):
    @staticmethod
    def quality_image(path: Path) -> None:
        image = Image.new("RGB", (1200, 700), "white")
        draw = ImageDraw.Draw(image)
        for value in range(0, 1200, 40):
            draw.rectangle((value, 100, value + 20, 650), fill=(value % 255, 80, 120))
        image.save(path)

    def test_windows_cp1252_child_still_emits_utf8(self):
        with tempfile.TemporaryDirectory(prefix="softcert-utf8-") as temp:
            project = Path(temp) / "中文 项目"
            project.mkdir()
            (project / "main.py").write_text("print('ok')\n", encoding="utf-8")
            env = os.environ.copy()
            env.update({"PYTHONUTF8": "0", "PYTHONIOENCODING": "cp1252"})
            result = subprocess.run(
                [sys.executable, str(SCRIPTS / "product_workflow.py"), "prepare", "--project", str(project)],
                capture_output=True, text=True, encoding="utf-8", errors="strict", env=env, timeout=60,
            )
            self.assertEqual(result.returncode, 3, result.stdout + result.stderr)
            self.assertNotIn("UnicodeEncodeError", result.stderr)

    def test_isolated_process_timeout_emits_diagnostic(self):
        with tempfile.TemporaryDirectory(prefix="softcert-timeout-") as temp:
            diagnostics = Path(temp)
            with self.assertRaises(ConversionFailure) as raised:
                run_isolated([sys.executable, "-c", "import time; time.sleep(30)"],
                             0.3, diagnostics, "timeout-probe")
            self.assertEqual(raised.exception.diagnostic["status"], "timeout")
            self.assertTrue((diagnostics / "timeout-probe.diagnostic.json").is_file())

    def test_manual_contains_materialized_clickable_toc(self):
        with tempfile.TemporaryDirectory(prefix="softcert-toc-") as temp:
            output = Path(temp) / "manual.docx"
            completed = subprocess.run([
                sys.executable, str(SCRIPTS / "build_manual.py"),
                "--input", str(ROOT / "assets/examples/manual-input.example.json"),
                "--theme", str(ROOT / "assets/themes/standard-filing-gray.json"),
                "--output", str(output),
            ], capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60)
            self.assertEqual(completed.returncode, 0, completed.stdout + completed.stderr)
            with zipfile.ZipFile(output) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")
                settings_xml = package.read("word/settings.xml").decode("utf-8")
                styles_xml = package.read("word/styles.xml").decode("utf-8")
            self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)
            self.assertIn('<w:hyperlink w:anchor="_SoftCertToc', document_xml)
            self.assertIn('<w:bookmarkStart', document_xml)
            self.assertIn('PAGEREF _SoftCertToc', document_xml)
            self.assertIn('<w:updateFields w:val="true"', settings_xml)
            self.assertIn('w:eastAsia="Microsoft YaHei"', styles_xml)
            self.assertNotIn('w:eastAsia="SimHei"', styles_xml)

    def test_step_blocks_restart_numbering_and_avoid_word_list_continuation(self):
        from docx import Document

        document = Document()
        theme = json.loads((ROOT / "assets/themes/standard-filing-gray.json").read_text(encoding="utf-8"))
        blocks = [
            {"type": "steps", "items": ["第一节步骤一", "第一节步骤二"]},
            {"type": "steps", "items": ["第二节步骤一", "第二节步骤二"]},
        ]
        add_blocks(document, blocks, theme, Path.cwd())
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs],
            ["1. 第一节步骤一", "2. 第一节步骤二", "1. 第二节步骤一", "2. 第二节步骤二"],
        )
        self.assertTrue(all(paragraph.style.name == "Normal" for paragraph in document.paragraphs))
    def test_application_rules_emit_original_source_count_as_digits(self):
        with tempfile.TemporaryDirectory(prefix="softcert-fields-") as temp:
            root = Path(temp)
            facts = {
                "software_full_name": "规则测试软件", "software_short_name": "规则测试", "version": "V1.0",
                "rightsholder": {"type": "legal_person", "name": "测试组织", "id_type": "测试标识", "id_number": "T-1"},
                "completion_date": "2026-08-11", "development_mode": "independent",
                "software_nature": "original", "publication": {"status": "unpublished"},
                "rights_acquisition": "original", "rights_scope": {"type": "all", "detail": ""},
                "ownership_notes": "",
            }
            analysis = {"technology": {"project_types": ["Web"]}, "field_inferences": {
                "programming_languages": {"suggested_value": ["Python"]},
                "development_tools": {"suggested_value": ["Python 3.12"]},
            }}
            business = {
                "software_classification": "应用软件", "software_purpose": "验证字段规则",
                "target_users": "测试员", "industry_domain": "软件服务", "development_environment": "Windows",
                "development_tools": "Python 3.12", "runtime_platform": "Windows", "runtime_support": "Python",
                "development_purpose": "验证外置规则的稳定性", "technical_features": "可追溯和确定性校验",
                "main_functions": "本软件通过实际项目信息完成字段提取、格式校验、条件必填判断和结果输出。" * 16,
                "capabilities": [],
            }
            provenance = {"original_line_count": 123, "full_line_count": 456}
            for name, value in (("facts", facts), ("analysis", analysis), ("business", business), ("provenance", provenance)):
                (root / f"{name}.json").write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            output, model = root / "application.txt", root / "model.json"
            completed = subprocess.run([
                sys.executable, str(SCRIPTS / "generate_application_form.py"),
                "--facts", str(root / "facts.json"), "--analysis", str(root / "analysis.json"),
                "--business", str(root / "business.json"), "--provenance", str(root / "provenance.json"),
                "--output", str(output), "--model-output", str(model),
            ], capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60)
            self.assertEqual(completed.returncode, 0, completed.stdout + completed.stderr)
            data = json.loads(model.read_text(encoding="utf-8"))
            source_field = next(item for item in data["fields"] if item["key"] == "source_line_count")
            self.assertEqual(source_field["value"], "123")
            self.assertIn("源程序量：123\n", output.read_text(encoding="utf-8"))
            self.assertNotIn("源程序量：123 行", output.read_text(encoding="utf-8"))
            self.assertTrue(data["live_limit_fields"])
            compressed, changed = compress("字" * 20, 10)
            self.assertTrue(changed)
            self.assertLessEqual(len(compressed), 10)

    def test_screenshot_states_are_not_vacuously_true(self):
        with tempfile.TemporaryDirectory(prefix="softcert-shots-") as temp:
            root = Path(temp)
            index = root / "screenshot-index.json"
            awaiting = {"mode": "chrome_devtools", "state": "awaiting_capture", "captures": [],
                        "summary": {"errors": 1, "missing_planned": 1}}
            index.write_text(json.dumps(awaiting), encoding="utf-8")
            self.assertFalse(screenshot_release_check(index, awaiting)[0])
            skipped = {"mode": "skip", "state": "skipped_by_user", "captures": [], "summary": {}}
            self.assertFalse(screenshot_release_check(index, skipped)[0])

            image_path = root / "shot.png"
            self.quality_image(image_path)
            captured = {"mode": "computer_use", "state": "captured", "captures": [{
                "id": "core", "status": "pass", "path": str(image_path), "sha256": sha256_file(image_path),
                "metrics": {"width": 1200, "height": 700}, "quality_findings": [],
                "evidence_ids": ["CAP-core"], "role": "测试角色", "url": "app://fixture/core",
            }], "summary": {"errors": 0, "quality_warnings": 0, "missing_planned": 0}}
            self.assertTrue(screenshot_release_check(index, captured)[0])

    def test_computer_use_receipt_is_finalized_to_release_index(self):
        with tempfile.TemporaryDirectory(prefix="softcert-computer-use-") as temp:
            root = Path(temp)
            source, plan, session = root / "source.png", root / "plan.json", root / "session.json"
            output, report = root / "screenshots", root / "screenshot-index.json"
            self.quality_image(source)
            plan.write_text(json.dumps({
                "schema_version": "1.0", "captures": [{
                    "id": "core-result", "title": "处理结果", "role": "测试员",
                    "route": "app://fixture/result", "evidence_ids": ["CAP-core"],
                }],
            }, ensure_ascii=False), encoding="utf-8")
            session.write_text(json.dumps({
                "started_at": "2026-08-11T00:00:00+08:00",
                "completed_at": "2026-08-11T00:01:00+08:00",
                "launch": {"result": "pass", "command": "fixture-app"},
                "login": {"required": True, "result": "pass", "role": "测试员"},
                "actions": [{"capture_id": "core-result", "action": "run", "result": "pass"}],
                "captures": [{"id": "core-result", "source_path": str(source),
                              "url": "app://fixture/result", "role": "测试员"}],
            }, ensure_ascii=False), encoding="utf-8")
            completed = subprocess.run([
                sys.executable, str(SCRIPTS / "finalize_agent_screenshots.py"),
                "--plan", str(plan), "--session", str(session), "--output", str(output),
                "--report", str(report),
            ], capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60)
            self.assertEqual(completed.returncode, 0, completed.stdout + completed.stderr)
            index = json.loads(report.read_text(encoding="utf-8"))
            self.assertEqual(index["state"], "captured")
            self.assertEqual(index["captures"][0]["id"], "core-result")
            self.assertTrue(Path(index["captures"][0]["path"]).is_file())

    def test_user_screenshot_keeps_planned_id_for_manual_mapping(self):
        with tempfile.TemporaryDirectory(prefix="softcert-user-shot-") as temp:
            root = Path(temp)
            source, output = root / "source", root / "output"
            source.mkdir()
            self.quality_image(source / "01-core.png")
            plan, report = root / "plan.json", root / "index.json"
            plan.write_text(json.dumps({"captures": [{
                "id": "planned-core", "title": "核心结果", "role": "操作员",
                "evidence_ids": ["CAP-core"],
            }]}, ensure_ascii=False), encoding="utf-8")
            completed = subprocess.run([
                sys.executable, str(SCRIPTS / "ingest_user_screenshots.py"),
                "--source", str(source), "--output", str(output), "--plan", str(plan),
                "--report", str(report),
            ], capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60)
            self.assertEqual(completed.returncode, 0, completed.stdout + completed.stderr)
            index = json.loads(report.read_text(encoding="utf-8"))
            self.assertEqual(index["state"], "captured")
            self.assertEqual(index["captures"][0]["id"], "planned-core")

    def test_manual_profiles_differ_and_repetition_is_detected(self):
        business = {"target_users": "业务人员"}
        base = {"purpose": "完成真实业务处理", "actor": "业务人员", "entry": "工作台",
                "visible_elements": "字段、按钮和结果区", "steps": ["输入条件", "执行处理", "核对结果"],
                "restrictions": ["仅处理当前权限范围"], "success_feedback": "显示处理成功",
                "error_feedback": "显示错误原因"}
        query = {**base, "name": "记录查询", "operation_type": "query"}
        approval = {**base, "name": "记录审批", "operation_type": "approval", "state_changes": ["待审变为通过"]}
        _, query_blocks, _ = capability_blocks(query, business, 1)
        _, approval_blocks, _ = capability_blocks(approval, business, 2)
        self.assertNotEqual(query_blocks[1]["headers"], approval_blocks[1]["headers"])
        repeated_page = {"title": "重复章节", "lead": "这是一个足够长的重复业务说明文本",
                         "blocks": [{"type": "paragraph", "text": "这是一个足够长的重复业务说明文本。"},
                                    {"type": "steps", "items": ["步骤一", "步骤二"]}]}
        quality = content_quality([repeated_page, repeated_page, repeated_page], {
            "minimum_capability_characters": 10, "minimum_steps_per_capability": 2,
            "minimum_detail_items": 1, "maximum_repeated_sentence_ratio": 0.1,
            "maximum_identical_block_signature_count": 2, "minimum_sentence_characters_for_repeat_check": 8,
        })
        self.assertEqual(quality["status"], "fail")

    def test_login_prerequisite_is_not_circular(self):
        capability = {"name": "账号登录", "entry": "/login",
                      "prerequisites": ["用户已登录系统", "访问地址可用"]}
        prerequisites = effective_prerequisites(capability)
        self.assertNotIn("用户已登录系统", prerequisites)
        self.assertIn("已取得系统分配的有效账号和认证信息", prerequisites)
        self.assertIn("访问地址可用", prerequisites)


if __name__ == "__main__":
    unittest.main()
